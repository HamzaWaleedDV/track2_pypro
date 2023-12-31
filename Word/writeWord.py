import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

mydoc = docx.Document()
mydoc.add_paragraph("This is first paragraph of a MS Word file.")
mydoc.add_paragraph("This is second paragraph of a MS Word file.")

paragraph = mydoc.add_paragraph(".هذه هي الفقرة الثالثة")
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

third_paragraph = mydoc.add_paragraph("This is the third paragraph.")
third_paragraph.add_run("this a section at the end of third paragraph.")

mydoc.add_heading("This is level 1 heading", 0)
mydoc.add_heading("This is level 2 heading", 1)
mydoc.add_heading("This is level 3 heading", 2)
mydoc.add_heading("This is level 4 heading", 3)
mydoc.add_heading("This is level 5 heading", 4)
mydoc.add_heading("This is level 6 heading", 5)

#style
print(mydoc.paragraphs[0].style)
print(mydoc.paragraphs[5].style)
print(mydoc.paragraphs[4].style)

mydoc.paragraphs[0].style = mydoc.styles['Heading 1']
mydoc.paragraphs[3].style = mydoc.styles['Heading 4']
mydoc.paragraphs[3].style.delete()

mydoc.add_paragraph("Hello, World", 'Title')

#add_photo
mydoc.add_picture(str(Path.home() / Path('OneDrive', 'Desktop', 'lll.png')), width = docx.shared.Inches(5), height = docx.shared.Inches(7))

mydoc.save(Path.home() / Path('OneDrive', 'Desktop', 'write.docx'))